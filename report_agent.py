"""
Report Generation Agent
리포트 생성 및 문서 저장
"""

from datetime import date, datetime
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json


def _source_map(sources: List[Dict[str, Any]]) -> Dict[str, List[int]]:
    """소스 인덱스 맵핑"""
    idxs = list(range(len(sources)))
    return {
        "claim_business": idxs[:3] or idxs,
        "claim_value": idxs[1:4] or idxs,
        "claim_tech": idxs[2:5] or idxs,
        "claim_risk": idxs[-3:] or idxs
    }


def _generate_recommendation(
    perspective: str,
    sections: Dict[str, Any],
    tables: Dict[str, Any],
    sources: List[Dict[str, Any]],
    openai_client
) -> str:
    """
    투자 추천 생성
    
    Args:
        perspective: 투자 관점 (value/tech)
        sections: 분석 섹션들
        tables: 데이터 테이블들
        sources: 출처 목록
        openai_client: OpenAI 클라이언트
    
    Returns:
        str: 투자 추천 텍스트 (Markdown)
    """
    sys = f"""Based on {perspective} perspective, generate investment recommendation (Buy/Hold/Sell).

Structure:
1. **Position**: Bold statement (e.g., **Buy**, **Hold**, **Sell**)
2. **Key Reasons**: 3-5 bullet points combining:
   - Our analysis: Summarize key metrics from sections and tables
   - Source insights: Reference analyst views and market consensus
3. **Target Price**: If available from sources
4. **Time Horizon**: Short/Medium/Long term view

Use Korean. Be concise (150-250 words). Focus on actionable insights."""

    user = json.dumps({
        "sections_summary": {
            k: v[:500] for k, v in sections.items() if v  # 요약만
        },
        "tables": {
            k: v for k, v in tables.items() if v
        },
        "top_sources": [
            {"domain": s.get("domain"), "title": s.get("title")} 
            for s in sources[:5]
        ],
        "perspective": perspective
    }, ensure_ascii=False)
    
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": sys},
                {"role": "user", "content": user}
            ],
            temperature=0.3
        )
        recommendation = resp.choices[0].message.content
        return "### 투자 추천\n\n" + recommendation
    except Exception as e:
        print(f"⚠️  Recommendation generation failed: {e}")
        return (
            "### 투자 추천\n\n"
            "**Hold (중립 유지)**\n\n"
            "- 추가 정보 수집 필요: 현재 데이터만으로는 명확한 방향성 판단이 어렵습니다.\n"
            "- 시장 모니터링: 주요 지표와 뉴스를 지속적으로 확인하시기 바랍니다.\n"
            "- 리스크 관리: 분산투자 원칙을 유지하는 것을 권장합니다.\n"
        )


def render(state: Dict[str, Any]) -> Tuple[str, str]:
    """
    Markdown 리포트 렌더링
    
    Args:
        state: 파이프라인 상태
    
    Returns:
        tuple: (report_md, tl_dr)
    """
    e = state["entity"]
    today = str(date.today())
    
    # 섹션 추출
    biz = state["sections"].get("business", "")
    val = state["sections"].get("value", "")
    tech = state["sections"].get("tech", "")
    risks = state["sections"].get("risks", "")
    
    openai_client = state["openai_client"]
    sources = state.get("sources", [])

    # TL;DR 생성
    parts = []
    if val:
        parts.append("가치")
    if tech:
        parts.append("기술")
    
    perspective_str = " · ".join(parts) if parts else "종합"
    tl = f"{e['name']}({e['ticker']}): 뉴스 및 공시 기반 {perspective_str} 관점 분석 ({today})"

    # Markdown 조합
    md = []
    md.append(f"# {e['name']} ({e['ticker']}) 투자 리포트\n")
    
    md.append("## TL;DR")
    md.append(f"- {tl}\n")
    
    md.append("## 기업 개요")
    md.append(f"- **티커**: {e['ticker']}")
    md.append(f"- **거래소**: {e.get('exchange', 'UNKNOWN')}")
    md.append(f"- **분석 경로**: {' → '.join(state.get('route', []))}")
    md.append(f"- **관점**: {state.get('perspective', 'value').upper()}\n")
    
    if biz:
        md.append(biz + "\n")
    
    if val:
        md.append(val + "\n")
    
    if tech:
        md.append(tech + "\n")
    
    if risks:
        md.append(risks + "\n")
    
    # 투자 추천 생성
    rec = _generate_recommendation(
        state["perspective"],
        state["sections"],
        state["tables"],
        sources,
        openai_client
    )
    md.append(rec + "\n")
    
    # 차트
    chart_path = state.get("artifacts", {}).get("price_chart")
    if chart_path:
        md.append(f"## 차트\n")
        md.append(f"![Price Chart]({chart_path})\n")
    
    # 출처
    if sources:
        md.append("## 출처\n")
        for i, s in enumerate(sources[:15], 1):  # 상위 15개
            domain = s.get('domain', '')
            title = s.get('title', 'Untitled')
            date_str = s.get('date', '')
            url = s.get('url', '')
            priority = s.get('priority', 0)
            
            # 우선순위 표시 (⭐)
            stars = "⭐" * min(priority // 2, 5) if priority > 5 else ""
            
            md.append(f"{i}. {stars} [{domain}] {title}")
            if date_str:
                md.append(f"   - Date: {date_str}")
            md.append(f"   - URL: {url}\n")
    
    # state 업데이트
    state["sections"]["recommendation"] = rec
    state["source_map"] = _source_map(sources)
    
    return "\n".join(md), tl


def _add_keyval_table(doc: Document, title: str, mapping: Dict[str, Any]) -> None:
    """Key-Value 테이블 추가"""
    if not mapping:
        return
    
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    
    try:
        table.style = "Light List"
    except:
        table.style = "Table Grid"
    
    hdr = table.rows[0].cells
    hdr[0].text = "항목"
    hdr[1].text = "값"
    
    # 헤더 볼드
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for k, v in mapping.items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(v) if v is not None else "N/A"


def _add_list_table(doc: Document, title: str, rows: List[Dict[str, Any]], columns: List[str]) -> None:
    """리스트 테이블 추가"""
    if not rows:
        return
    
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(columns))
    
    try:
        table.style = "Light List"
    except:
        table.style = "Table Grid"
    
    # 헤더
    for i, c in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = c
        cell.paragraphs[0].runs[0].bold = True
    
    # 데이터
    for r in rows:
        row = table.add_row().cells
        for i, c in enumerate(columns):
            row[i].text = str(r.get(c, ""))


def save_docx(state: Dict[str, Any], out_path: Optional[str] = None) -> str:
    """
    Word 문서 저장
    
    Args:
        state: 파이프라인 상태
        out_path: 출력 경로 (선택)
    
    Returns:
        str: 저장된 파일 경로
    """
    e = state.get("entity", {}) or {}
    doc = Document()
    
    # 제목
    title = doc.add_heading(f"{e.get('name', '기업')} ({e.get('ticker', '?')}) 투자 리포트", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # TL;DR
    doc.add_heading("TL;DR", level=1)
    p = doc.add_paragraph(state.get("tl_dr", ""))
    p.runs[0].font.size = Pt(11)
    
    # 기업 개요
    doc.add_heading("기업 개요", level=1)
    overview = doc.add_paragraph()
    overview.add_run(f"티커: ").bold = True
    overview.add_run(f"{e.get('ticker', '?')}\n")
    overview.add_run(f"거래소: ").bold = True
    overview.add_run(f"{e.get('exchange', 'UNKNOWN')}\n")
    overview.add_run(f"분석 경로: ").bold = True
    overview.add_run(f"{', '.join(state.get('route', []))}\n")
    overview.add_run(f"관점: ").bold = True
    overview.add_run(f"{state.get('perspective', 'value').upper()}\n")
    
    # 섹션들
    sections = [
        ("사업 전개 및 동향", state.get("sections", {}).get("business", "")),
        ("가치투자 분석", state.get("sections", {}).get("value", "")),
        ("기술적 분석", state.get("sections", {}).get("tech", "")),
        ("리스크 분석", state.get("sections", {}).get("risks", "")),
        ("투자 추천", state.get("sections", {}).get("recommendation", ""))
    ]
    
    for section_title, body in sections:
        if body:
            doc.add_heading(section_title, level=1)
            
            # Markdown 스타일 파싱 (간단한 버전)
            lines = body.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 헤딩 제거 (이미 추가됨)
                if line.startswith('#'):
                    continue
                
                p = doc.add_paragraph()
                
                # 볼드 처리
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 0:  # 일반 텍스트
                        p.add_run(part)
                    else:  # 볼드
                        p.add_run(part).bold = True
    
    # 테이블들
    t = state.get("tables", {}) or {}
    
    if "rev_compare" in t and t["rev_compare"]:
        _add_list_table(doc, "분기 매출 비교", t["rev_compare"], ["Quarter", "Last", "Prev/YoY", "Type"])
    
    fb = t.get("finviz_blocks", {})
    if fb.get("valuation"):
        _add_keyval_table(doc, "밸류에이션 멀티플", fb["valuation"])
    if fb.get("profit_leverage"):
        _add_keyval_table(doc, "수익성 & 레버리지", fb["profit_leverage"])
    
    if "technicals" in t:
        _add_keyval_table(doc, "기술 지표 (최근)", t["technicals"])
    if "trend" in t:
        _add_keyval_table(doc, "추세 분석 (60일)", t["trend"])
    
    if "filings" in t and t["filings"]:
        _add_list_table(doc, "SEC 공시 (최근)", t["filings"], ["Form", "Date", "URL"])
    
    # 차트
    chart_path = state.get("artifacts", {}).get("price_chart")
    if chart_path and Path(chart_path).exists():
        doc.add_heading("차트", level=1)
        try:
            doc.add_picture(chart_path, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"차트를 로드할 수 없습니다: {e}")
    
    # 출처
    sources = state.get("sources", [])
    if sources:
        doc.add_heading("출처", level=1)
        for i, s in enumerate(sources[:15], 1):
            p = doc.add_paragraph()
            
            # 우선순위 표시
            priority = s.get("priority", 0)
            stars = "⭐" * min(priority // 2, 5) if priority > 5 else ""
            
            p.add_run(f"{i}. {stars} ").bold = True
            p.add_run(f"[{s.get('domain', '')}] {s.get('title', '')}")
            
            if s.get('date'):
                p.add_run(f"\n   Date: {s['date']}")
            p.add_run(f"\n   URL: {s.get('url', '')}\n")
    
    # 파일 저장
    out_dir = Path("reports")
    out_dir.mkdir(parents=True, exist_ok=True)
    
    if out_path:
        path = Path(out_path)
    else:
        ticker = e.get('ticker', 'REPORT')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        path = out_dir / f"{ticker}_{timestamp}.docx"
    
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    
    doc.save(str(path))
    return str(path)
